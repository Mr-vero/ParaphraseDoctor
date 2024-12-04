import docx
from pypdf import PdfReader
from transformers import AutoTokenizer, AutoModelForSeq2SeqLM
import torch
from docx import Document
import os
from typing import Callable
import gradio as gr
from tqdm import tqdm
import re
from concurrent.futures import ThreadPoolExecutor
import numpy as np
from docx.shared import Pt, RGBColor
import math
from google.generativeai import GenerativeModel
import google.generativeai as genai
from dotenv import load_dotenv
import logging
from datetime import datetime
import json
import time

class DocumentParaphraser:
    def __init__(self):
        load_dotenv()  # Load environment variables from .env file
        api_key = os.getenv('GEMINI_API_KEY')
        if not api_key:
            raise ValueError("GEMINI_API_KEY not found in environment variables")
        
        genai.configure(api_key=api_key)
        
        # Configure optimized generation settings
        generation_config = {
            "temperature": 1,
            "top_p": 0.95,
            "top_k": 40,
            "max_output_tokens": 8192,
            # "response_mime_type": "text/plain",
        }
        
        self.model = genai.GenerativeModel(
            model_name="gemini-1.5-flash",
            generation_config=generation_config
        )
        
        # Initialize chat session for better context handling
        self.chat_session = self.model.start_chat(history=[])
        
        # Set up logging
        log_dir = "logs"
        if not os.path.exists(log_dir):
            os.makedirs(log_dir)
            
        log_file = f"logs/llm_{datetime.now().strftime('%Y%m%d_%H%M%S')}.log"
        logging.basicConfig(
            level=logging.INFO,
            format='%(asctime)s - %(levelname)s - %(message)s',
            handlers=[
                logging.FileHandler(log_file),
                logging.StreamHandler()  # This will also print to console
            ]
        )
        self.logger = logging.getLogger(__name__)
        self.log_messages = []  # Add this line to store log messages
          # Initialize rate limiting
        self.usage_state_file = "gemini_usage_state.json"
        self.usage_state = self._load_usage_state()
        
    def get_usage_stats(self):
        """Return current usage statistics"""
        return {
            "requests_today": self.usage_state["requests"],
            "requests_remaining": 1500 - self.usage_state["requests"],
            "tokens_used": self.usage_state["tokens"]
        }
        
    def _load_usage_state(self):
        """Load usage state from a file."""
        if os.path.exists(self.usage_state_file):
            with open(self.usage_state_file, 'r') as file:
                return json.load(file)
        else:
            return {"requests": 0, "tokens": 0}

    def _save_usage_state(self):
        """Save usage state to a file."""
        with open(self.usage_state_file, 'w') as file:
            json.dump(self.usage_state, file)

    def _paraphrase_chunk(self, text):
        max_retries = 3
        retry_delay = 60  # 60 seconds delay when hitting rate limit
        
        for attempt in range(max_retries):
            try:
                # Check rate limits
                if self.usage_state["requests"] >= 1500:
                    raise Exception("Daily request limit reached.")
                if self.usage_state["requests"] >= 15:
                    log_msg = "Minute request limit reached. Waiting 60 seconds..."
                    self.logger.warning(log_msg)
                    self.log_messages.append(log_msg)  # Store the log message
                    time.sleep(retry_delay)
                    # Reset minute counter after waiting
                    self.usage_state["requests"] = 0
                    self._save_usage_state()
                
                prompt = f"""Paraphrase the following text while:
                - Maintaining a formal academic tone suitable for scientific journals or Indonesian "skripsi"
                - DO NOT modify any of the following:
                    * Citations (e.g., [Author, Year])
                    * Footnotes
                    * Religious texts/quotes (from Quran, Bible, or other religious sources)
                - Preserving all technical terms
                - Keeping the same meaning
                - IMPORTANT: Keep the exact same language as the input - DO NOT TRANSLATE
                - If the input is not in English, paraphrase in that same language
                
                Text: {text}
                
                Remember: 
                - Your paraphrase must be in the SAME LANGUAGE as the input text
                - Keep all citations, footnotes, and religious quotes exactly as they appear in the original text"""
                
                # Log the prompt
                log_msg = f"Sending prompt to LLM:\n{prompt}"
                self.logger.info(log_msg)
                self.log_messages.append(log_msg)  # Store the log message
                
                response = self.chat_session.send_message(prompt)
                # Log the response
                log_msg = f"Received response from LLM:\n{response.text}"
                self.logger.info(log_msg)
                self.log_messages.append(log_msg)  # Store the log message

                # Update usage state
                self.usage_state["requests"] += 1
                self.usage_state["tokens"] += len(response.text.split())
                self._save_usage_state()

                return response.text
                
            except Exception as e:
                if "rate limit" in str(e).lower() and attempt < max_retries - 1:
                    wait_time = retry_delay * (attempt + 1)
                    self.logger.warning(f"Rate limit hit. Waiting {wait_time} seconds before retry...")
                    time.sleep(wait_time)
                    continue
                else:
                    self.logger.error(f"Error in Gemini API call: {str(e)}")
                    raise
    
    def paraphrase_text(self, text_structure, progress_callback=None):
        paraphrased_structure = []
        total_items = len(text_structure)
        
        for idx, item in enumerate(text_structure):
            if progress_callback:
                progress_callback(f"Processing section {idx + 1}/{total_items}", idx/total_items)
            
            # Skip headings
            if item['type'] == 'heading':
                paraphrased_structure.append(item.copy())
                continue
            
            # Paraphrase content
            paraphrased_item = item.copy()
            paraphrased_item['text'] = self._paraphrase_chunk(item['text'])
            paraphrased_structure.append(paraphrased_item)
        
        return paraphrased_structure

    def _save_docx(self, structured_text, output_path):
        """Enhanced DOCX saving with better formatting"""
        doc = Document()
        
        # Define styles
        styles = {
            'Heading1': {'size': 16, 'bold': True},
            'Heading2': {'size': 14, 'bold': True},
            'Heading3': {'size': 12, 'bold': True},
            'Normal': {'size': 11, 'bold': False}
        }
        
        for item in structured_text:
            if item['type'] == 'heading':
                heading = doc.add_heading(item['text'], level=int(item['level']))
                heading.alignment = item['alignment']
                
                # Apply heading style
                style_name = f"Heading{item['level']}"
                if style_name in styles:
                    for run in heading.runs:
                        run.font.size = Pt(styles[style_name]['size'])
                        run.font.bold = styles[style_name]['bold']
            else:
                paragraph = doc.add_paragraph()
                run = paragraph.add_run(item['text'])
                paragraph.style = item['style']
                paragraph.alignment = item['alignment']
                
                # Apply normal style
                run.font.size = Pt(styles['Normal']['size'])
                
                # Handle special content formatting
                if '[' in item['text'] and ']' in item['text']:
                    # Format citations differently
                    citations = re.findall(self.patterns['citation'], item['text'])
                    for citation in citations:
                        text_parts = item['text'].split(f'[{citation}]')
                        for i, part in enumerate(text_parts):
                            run = paragraph.add_run(part)
                            if i < len(text_parts) - 1:
                                citation_run = paragraph.add_run(f'[{citation}]')
                                citation_run.font.color.rgb = RGBColor(0, 0, 255)
        
        doc.save(output_path)

    def _split_text_carefully(self, text, max_length):
        """Split text into smaller chunks while preserving sentence and context"""
        import re
        
        # Split by sentences while preserving punctuation
        sentences = re.split(r'(?<=[.!?])\s+', text)
        chunks = []
        current_chunk = []
        current_length = 0
        
        for sentence in sentences:
            sentence = sentence.strip()
            if not sentence:
                continue
                
            sentence_length = len(sentence)
            
            if current_length + sentence_length <= max_length:
                current_chunk.append(sentence)
                current_length += sentence_length
            else:
                if current_chunk:
                    chunks.append(' '.join(current_chunk))
                current_chunk = [sentence]
                current_length = sentence_length
        
        if current_chunk:
            chunks.append(' '.join(current_chunk))
        
        return chunks

    def _post_process_chunk(self, text):
        """Clean up and verify paraphrased text"""
        # Remove duplicate spaces
        text = ' '.join(text.split())
        # Ensure proper capitalization
        text = text[0].upper() + text[1:] if text else text
        # Ensure proper punctuation
        if text and text[-1] not in '.!?':
            text += '.'
        return text

    def _combine_chunks_with_citations(self, chunks, citations):
        """Combine paraphrased chunks while restoring citations"""
        text = ' '.join(chunks)
        
        # Restore citations if they exist
        if citations:
            for citation in citations:
                # Add citation at appropriate positions
                text = text.rstrip('.') + f' [{citation}].'
        
        return text

    def save_document(self, text, output_path):
        """Save the paraphrased text to a document"""
        file_extension = os.path.splitext(output_path)[1].lower()
        
        if file_extension == '.docx':
            self._save_docx(text, output_path)
        elif file_extension == '.txt':
            self._save_txt(text, output_path)
        else:
            raise ValueError(f"Unsupported output format: {file_extension}")

    def _save_txt(self, structured_text, output_path):
        """Save text as TXT with basic formatting"""
        with open(output_path, 'w', encoding='utf-8') as file:
            for item in structured_text:
                if item['type'] == 'heading':
                    file.write('\n' + '=' * 40 + '\n')
                    file.write(item['text'].upper() + '\n')
                    file.write('=' * 40 + '\n\n')
                else:
                    file.write(item['text'] + '\n\n')

    def process_document(self, input_path, output_path, progress_callback=None):
        """Process entire document with structure preservation"""
        try:
            # Verify input file exists
            if not os.path.exists(input_path):
                raise FileNotFoundError(f"Input file not found: {input_path}")

            # Read document with structure
            if progress_callback:
                progress_callback("Reading document...", 0.1)
            
            # Determine file type and use appropriate reader
            file_extension = os.path.splitext(input_path)[1].lower()
            if file_extension == '.txt':
                document_structure = self._read_txt(input_path)
            elif file_extension == '.docx':
                document_structure = self._read_docx(input_path)
            elif file_extension == '.pdf':
                document_structure = self._read_pdf(input_path)
            else:
                raise ValueError(f"Unsupported file format: {file_extension}")
            
            try:
                paraphrased_structure = self.paraphrase_text(
                    document_structure,
                    progress_callback=progress_callback
                )
            except Exception as e:
                self.logger.error(f"Paraphrasing error: {str(e)}")
                raise Exception(f"An error occurred during paraphrasing: {str(e)}")
            
            # Save with preserved structure
            if progress_callback:
                progress_callback("Saving document...", 0.9)
            
            # Use appropriate save method based on output format
            output_extension = os.path.splitext(output_path)[1].lower()
            if output_extension == '.txt':
                self._save_txt(paraphrased_structure, output_path)
            elif output_extension == '.docx':
                self._save_docx(paraphrased_structure, output_path)
            else:
                raise ValueError(f"Unsupported output format: {output_extension}")
            
            return output_path
            
        except Exception as e:
            self.logger.error(f"Document processing error: {str(e)}")
            raise Exception(f"An error occurred: {str(e)}")

    def _read_docx(self, file_path):
        """Extract text from DOCX with research paper structure preservation"""
        doc = Document(file_path)
        document_structure = []
        
        for paragraph in doc.paragraphs:
            style = paragraph.style.name
            text = paragraph.text.strip()
            
            if not text:  # Skip empty paragraphs
                continue
            
            # Create structured content
            content = {
                'type': 'paragraph' if not style.startswith('Heading') else 'heading',
                'style': style,
                'text': text,
                'alignment': paragraph.alignment
            }
            
            # Add level for headings
            if content['type'] == 'heading':
                try:
                    content['level'] = int(style[-1])
                except (ValueError, IndexError):
                    content['level'] = 1
                
            document_structure.append(content)
        
        return document_structure

    def _read_pdf(self, file_path):
        """Extract text from PDF with structure preservation"""
        reader = PdfReader(file_path)
        document_structure = []
        
        for page in reader.pages:
            text = page.extract_text()
            if text.strip():
                document_structure.append({
                    'type': 'paragraph',
                    'style': 'Normal',
                    'text': text,
                    'alignment': 0  # Left alignment as default
                })
        
        return document_structure

    def _read_txt(self, file_path):
        """Read text file with basic structure preservation"""
        document_structure = []
        
        with open(file_path, 'r', encoding='utf-8') as file:
            for line in file:
                text = line.strip()
                if text:
                    document_structure.append({
                        'type': 'paragraph',
                        'style': 'Normal',
                        'text': text,
                        'alignment': 0  # Left alignment as default
                    })
        
        return document_structure

    def read_document(self, file_path):
        """Read different document formats and extract text with structure"""
        file_extension = os.path.splitext(file_path)[1].lower()
        
        if file_extension == '.pdf':
            return self._read_pdf(file_path)
        elif file_extension == '.docx':
            return self._read_docx(file_path)
        elif file_extension == '.txt':
            return self._read_txt(file_path)
        else:
            raise ValueError(f"Unsupported file format: {file_extension}")